import streamlit as st
from google import genai
from google.genai import types
import tempfile
import os
import io
import re
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =========================
# 1. CẤU HÌNH CHUNG
# =========================
st.set_page_config(page_title="Trợ lý Giáo án NLS", page_icon="📘", layout="centered")

FILE_KHUNG_NANG_LUC = "khungnanglucso.pdf"
MODEL_NAME = "gemini-2.5-flash-lite"

TEN_TAC_GIA = "La Mạnh Hùng"
TEN_TRUONG = "Trường PTDTBT TH&THCS Nà Khương"
SO_DIEN_THOAI = "0388 667 404"
TEN_UNG_DUNG = "📘 TRỢ LÝ SOẠN GIÁO ÁN TỰ ĐỘNG (NLS)"


# =========================
# 2. HÀM TIỆN ÍCH
# =========================
def safe_html(text: str) -> str:
    if not text:
        return ""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def clean_text(s: str) -> str:
    if not s:
        return ""
    s = str(s).strip()
    banned_prefixes = [
        "tuyệt vời",
        "hy vọng",
        "nếu có bất kỳ",
        "tôi rất sẵn lòng",
        "dựa vào khung chương trình",
        "xin chào",
        "sau đây",
        "mình sẽ",
        "tôi sẽ",
    ]
    lower = s.lower()
    for bad in banned_prefixes:
        if lower.startswith(bad):
            return ""
    return s


def normalize_bullets(items):
    result = []
    if not items:
        return result
    for item in items:
        item = clean_text(item)
        if item:
            result.append(item)
    return result


def set_run_font(run, bold=False, size=14):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_bullet_paragraph(doc_or_cell, text):
    p = doc_or_cell.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, bold=False, size=14)
    return p


def add_normal_paragraph(doc_or_cell, text, bold=False, align=None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, bold=bold, size=14)
    return p


def extract_json_block(text: str) -> dict:
    if not text:
        raise ValueError("Model không trả về nội dung.")

    text = text.strip()
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("Không tìm thấy JSON hợp lệ trong phản hồi của AI.")

    json_text = text[start:end + 1]
    return json.loads(json_text)


def validate_lesson_plan(data: dict):
    required_top_keys = [
        "ten_bai",
        "mon_hoc",
        "lop",
        "thoi_luong",
        "yeu_cau_can_dat",
        "do_dung_day_hoc",
        "tien_trinh_day_hoc",
        "dieu_chinh_sau_tiet_day"
    ]
    for key in required_top_keys:
        if key not in data:
            raise ValueError(f"Thiếu trường bắt buộc trong JSON: {key}")

    yccd = data["yeu_cau_can_dat"]
    for key in [
        "hoc_sinh_thuc_hien_duoc",
        "hoc_sinh_van_dung_duoc",
        "phat_trien_nang_luc",
        "phat_trien_pham_chat",
        "noi_dung_tich_hop"
    ]:
        if key not in yccd:
            raise ValueError(f"Thiếu mục trong 'yeu_cau_can_dat': {key}")

    ptnl = yccd["phat_trien_nang_luc"]
    for key in ["nang_luc_dac_thu", "nang_luc_chung", "nang_luc_so"]:
        if key not in ptnl:
            raise ValueError(f"Thiếu mục trong 'phat_trien_nang_luc': {key}")

    ndth = yccd["noi_dung_tich_hop"]
    for key in ["hoc_thong_qua_choi", "cong_dan_so"]:
        if key not in ndth:
            raise ValueError(f"Thiếu mục trong 'noi_dung_tich_hop': {key}")

    dddh = data["do_dung_day_hoc"]
    for key in ["giao_vien", "hoc_sinh"]:
        if key not in dddh:
            raise ValueError(f"Thiếu mục trong 'do_dung_day_hoc': {key}")

    activities = data["tien_trinh_day_hoc"]
    if not isinstance(activities, list) or len(activities) != 4:
        raise ValueError("Tiến trình dạy học phải có đúng 4 hoạt động.")

    expected_names = [
        "Hoạt động 1 - Khởi động",
        "Hoạt động 2 - Hình thành kiến thức mới",
        "Hoạt động 3 - Thực hành - luyện tập",
        "Hoạt động 4 - Vận dụng"
    ]

    for idx, act in enumerate(activities):
        for key in ["ten_hoat_dong", "thoi_gian", "giao_vien", "hoc_sinh"]:
            if key not in act:
                raise ValueError(f"Thiếu trường '{key}' trong hoạt động {idx + 1}")

        if expected_names[idx].lower() not in act["ten_hoat_dong"].lower():
            raise ValueError(f"Tên hoạt động {idx + 1} chưa đúng chuẩn.")

        if not isinstance(act["giao_vien"], list) or not act["giao_vien"]:
            raise ValueError(f"Hoạt động {idx + 1} thiếu nội dung cột giáo viên.")

        if not isinstance(act["hoc_sinh"], list) or not act["hoc_sinh"]:
            raise ValueError(f"Hoạt động {idx + 1} thiếu nội dung cột học sinh.")


def ensure_default_adjustments(data: dict):
    if not data.get("dieu_chinh_sau_tiet_day"):
        data["dieu_chinh_sau_tiet_day"] = [
            "Sau tiết dạy, giáo viên ghi nhận những điểm phù hợp và những nội dung cần điều chỉnh để tổ chức bài học hiệu quả hơn ở lần sau."
        ]
    return data


def upload_file_to_gemini(client, file_path, mime_type):
    return client.files.upload(
        file=file_path,
        config=types.UploadFileConfig(mime_type=mime_type)
    )


def build_contents(prompt_instruction, uploaded_refs, extra_text=None):
    parts = [types.Part(text=prompt_instruction)]

    for ref in uploaded_refs:
        parts.append(types.Part.from_uri(file_uri=ref.uri, mime_type=ref.mime_type))

    if extra_text:
        parts.append(types.Part(text=extra_text))

    return [types.Content(role="user", parts=parts)]


def render_lesson_plan_text(data: dict) -> str:
    lines = []
    lines.append(f"KẾ HOẠCH BÀI DẠY: {data.get('ten_bai', '').upper()}")
    lines.append(f"Môn học: {data.get('mon_hoc', '')}")
    lines.append(f"Lớp: {data.get('lop', '')}")
    lines.append(f"Thời lượng: {data.get('thoi_luong', '35 phút')}")
    lines.append("")

    lines.append("I. Yêu cầu cần đạt")
    yccd = data["yeu_cau_can_dat"]

    lines.append("1. Học sinh thực hiện được")
    for x in normalize_bullets(yccd["hoc_sinh_thuc_hien_duoc"]):
        lines.append(f"- {x}")

    lines.append("2. Học sinh vận dụng được")
    for x in normalize_bullets(yccd["hoc_sinh_van_dung_duoc"]):
        lines.append(f"- {x}")

    lines.append("3. Phát triển năng lực")
    lines.append("- Năng lực đặc thù")
    for x in normalize_bullets(yccd["phat_trien_nang_luc"]["nang_luc_dac_thu"]):
        lines.append(f"  - {x}")

    lines.append("- Năng lực chung")
    for x in normalize_bullets(yccd["phat_trien_nang_luc"]["nang_luc_chung"]):
        lines.append(f"  - {x}")

    lines.append("- Năng lực số")
    for x in normalize_bullets(yccd["phat_trien_nang_luc"]["nang_luc_so"]):
        lines.append(f"  - {x}")

    lines.append("4. Phát triển phẩm chất")
    for x in normalize_bullets(yccd["phat_trien_pham_chat"]):
        lines.append(f"- {x}")

    lines.append("* Nội dung tích hợp")
    lines.append("- Học thông qua chơi")
    for x in normalize_bullets(yccd["noi_dung_tich_hop"]["hoc_thong_qua_choi"]):
        lines.append(f"  - {x}")

    lines.append("- Công dân số")
    for x in normalize_bullets(yccd["noi_dung_tich_hop"]["cong_dan_so"]):
        lines.append(f"  - {x}")

    lines.append("")
    lines.append("II. Đồ dùng dạy học")
    lines.append("1. Giáo viên")
    for x in normalize_bullets(data["do_dung_day_hoc"]["giao_vien"]):
        lines.append(f"- {x}")

    lines.append("2. Học sinh")
    for x in normalize_bullets(data["do_dung_day_hoc"]["hoc_sinh"]):
        lines.append(f"- {x}")

    lines.append("")
    lines.append("III. Tiến trình dạy học")
    lines.append("| HOẠT ĐỘNG CỦA GIÁO VIÊN | HOẠT ĐỘNG CỦA HỌC SINH |")
    lines.append("|---|---|")
    for act in data["tien_trinh_day_hoc"]:
        gv_text = "<br>".join([f"- {x}" for x in normalize_bullets(act["giao_vien"])])
        hs_text = "<br>".join([f"- {x}" for x in normalize_bullets(act["hoc_sinh"])])
        title = f"**{act['ten_hoat_dong']} ({act['thoi_gian']})**<br>"
        lines.append(f"| {title}{gv_text} | {hs_text} |")

    lines.append("")
    lines.append("IV. Điều chỉnh sau tiết dạy")
    for x in normalize_bullets(data["dieu_chinh_sau_tiet_day"]):
        lines.append(f"- {x}")

    return "\n".join(lines)


def render_response_box(text: str):
    escaped = safe_html(text).replace("\n", "<br>")
    st.markdown("### 📄 KẾT QUẢ BÀI SOẠN:")
    st.markdown(
        f'<div class="lesson-plan-paper">{escaped}</div>',
        unsafe_allow_html=True
    )


def create_doc_from_json(data: dict):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.2

    # Tiêu đề
    head = doc.add_heading(f"KẾ HOẠCH BÀI DẠY: {data.get('ten_bai', '').upper()}", 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in head.runs:
        set_run_font(run, bold=True, size=14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Thông tin đầu
    add_normal_paragraph(doc, f"Môn học: {data.get('mon_hoc', '')}")
    add_normal_paragraph(doc, f"Lớp: {data.get('lop', '')}")
    add_normal_paragraph(doc, f"Thời lượng: {data.get('thoi_luong', '35 phút')}")

    # I
    add_normal_paragraph(doc, "I. Yêu cầu cần đạt", bold=True)
    yccd = data["yeu_cau_can_dat"]

    add_normal_paragraph(doc, "1. Học sinh thực hiện được", bold=True)
    for item in normalize_bullets(yccd["hoc_sinh_thuc_hien_duoc"]):
        add_bullet_paragraph(doc, item)

    add_normal_paragraph(doc, "2. Học sinh vận dụng được", bold=True)
    for item in normalize_bullets(yccd["hoc_sinh_van_dung_duoc"]):
        add_bullet_paragraph(doc, item)

    add_normal_paragraph(doc, "3. Phát triển năng lực", bold=True)

    add_normal_paragraph(doc, "- Năng lực đặc thù", bold=True)
    for item in normalize_bullets(yccd["phat_trien_nang_luc"]["nang_luc_dac_thu"]):
        add_bullet_paragraph(doc, item)

    add_normal_paragraph(doc, "- Năng lực chung", bold=True)
    for item in normalize_bullets(yccd["phat_trien_nang_luc"]["nang_luc_chung"]):
        add_bullet_paragraph(doc, item)

    add_normal_paragraph(doc, "- Năng lực số", bold=True)
    for item in normalize_bullets(yccd["phat_trien_nang_luc"]["nang_luc_so"]):
        add_bullet_paragraph(doc, item)

    add_normal_paragraph(doc, "4. Phát triển phẩm chất", bold=True)
    for item in normalize_bullets(yccd["phat_trien_pham_chat"]):
        add_bullet_paragraph(doc, item)

    add_normal_paragraph(doc, "* Nội dung tích hợp", bold=True)

    add_normal_paragraph(doc, "- Học thông qua chơi", bold=True)
    for item in normalize_bullets(yccd["noi_dung_tich_hop"]["hoc_thong_qua_choi"]):
        add_bullet_paragraph(doc, item)

    add_normal_paragraph(doc, "- Công dân số", bold=True)
    for item in normalize_bullets(yccd["noi_dung_tich_hop"]["cong_dan_so"]):
        add_bullet_paragraph(doc, item)

    # II
    add_normal_paragraph(doc, "II. Đồ dùng dạy học", bold=True)

    add_normal_paragraph(doc, "1. Giáo viên", bold=True)
    for item in normalize_bullets(data["do_dung_day_hoc"]["giao_vien"]):
        add_bullet_paragraph(doc, item)

    add_normal_paragraph(doc, "2. Học sinh", bold=True)
    for item in normalize_bullets(data["do_dung_day_hoc"]["hoc_sinh"]):
        add_bullet_paragraph(doc, item)

    # III
    add_normal_paragraph(doc, "III. Tiến trình dạy học", bold=True)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = True

    hdr = table.rows[0].cells
    hdr[0].text = "HOẠT ĐỘNG CỦA GIÁO VIÊN"
    hdr[1].text = "HOẠT ĐỘNG CỦA HỌC SINH"

    for cell in hdr:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True, size=14)

    for act in data["tien_trinh_day_hoc"]:
        row = table.add_row().cells

        # Cột giáo viên
        row[0].paragraphs[0].clear()
        p_left_title = row[0].paragraphs[0]
        p_left_title.paragraph_format.space_after = Pt(3)
        r_left = p_left_title.add_run(f"{act['ten_hoat_dong']} ({act['thoi_gian']})")
        set_run_font(r_left, bold=True, size=14)

        for item in normalize_bullets(act["giao_vien"]):
            add_bullet_paragraph(row[0], item)

        # Cột học sinh
        row[1].paragraphs[0].clear()
        p_right_title = row[1].paragraphs[0]
        p_right_title.paragraph_format.space_after = Pt(3)
        r_right = p_right_title.add_run("Hoạt động của học sinh")
        set_run_font(r_right, bold=True, size=14)

        for item in normalize_bullets(act["hoc_sinh"]):
            add_bullet_paragraph(row[1], item)

    # IV
    add_normal_paragraph(doc, "IV. Điều chỉnh sau tiết dạy", bold=True)
    for item in normalize_bullets(data["dieu_chinh_sau_tiet_day"]):
        add_bullet_paragraph(doc, item)

    return doc


def generate_lesson_prompt(ten_bai, lop, noidung_bosung, yeu_cau_them):
    return f"""
Bạn là công cụ sinh giáo án tiểu học theo đúng mẫu quy định.
NHIỆM VỤ: Sinh duy nhất 1 JSON hợp lệ. KHÔNG viết lời mở đầu. KHÔNG giải thích. KHÔNG kết luận. KHÔNG dùng markdown. KHÔNG dùng code fence.

Dữ liệu bài dạy:
- Tên bài: {ten_bai}
- Lớp: {lop}
- Ghi chú bổ sung: {noidung_bosung}
- Yêu cầu thêm: {yeu_cau_them}

Yêu cầu bắt buộc:
1. Bám đúng cấu trúc:
   I. Yêu cầu cần đạt
   II. Đồ dùng dạy học
   III. Tiến trình dạy học
   IV. Điều chỉnh sau tiết dạy

2. Mục I phải gồm đủ:
   - hoc_sinh_thuc_hien_duoc
   - hoc_sinh_van_dung_duoc
   - phat_trien_nang_luc:
       + nang_luc_dac_thu
       + nang_luc_chung
       + nang_luc_so
   - phat_trien_pham_chat
   - noi_dung_tich_hop:
       + hoc_thong_qua_choi
       + cong_dan_so

3. Mục III phải có đúng 4 hoạt động:
   - Hoạt động 1 - Khởi động
   - Hoạt động 2 - Hình thành kiến thức mới
   - Hoạt động 3 - Thực hành - luyện tập
   - Hoạt động 4 - Vận dụng

4. Mỗi hoạt động phải có:
   - ten_hoat_dong
   - thoi_gian
   - giao_vien: danh sách các ý, mỗi ý là 1 gạch đầu dòng
   - hoc_sinh: danh sách các ý, mỗi ý là 1 gạch đầu dòng

5. Phần III bắt buộc để app xuất thành bảng 2 cột:
   - Cột 1: HOẠT ĐỘNG CỦA GIÁO VIÊN
   - Cột 2: HOẠT ĐỘNG CỦA HỌC SINH
   Không được trộn lẫn hai cột.

6. Nếu có trò chơi thì luật chơi phải nằm trong danh sách 'giao_vien' của hoạt động tương ứng.

7. Tổng thời lượng phù hợp 35 phút.

8. Không được sinh các câu xã giao như:
   - Tuyệt vời...
   - Tôi rất sẵn lòng...
   - Hy vọng...
   - Nếu cần...
   - Đừng ngần ngại...

9. Mọi nội dung phải ngắn gọn, đúng chất giáo án, bám sát chuẩn trường tiểu học Việt Nam.

10. Chỉ trả về JSON hợp lệ theo đúng schema dưới đây:

{{
  "ten_bai": "{ten_bai}",
  "mon_hoc": "...",
  "lop": "{lop}",
  "thoi_luong": "35 phút",
  "yeu_cau_can_dat": {{
    "hoc_sinh_thuc_hien_duoc": ["..."],
    "hoc_sinh_van_dung_duoc": ["..."],
    "phat_trien_nang_luc": {{
      "nang_luc_dac_thu": ["..."],
      "nang_luc_chung": ["..."],
      "nang_luc_so": ["..."]
    }},
    "phat_trien_pham_chat": ["..."],
    "noi_dung_tich_hop": {{
      "hoc_thong_qua_choi": ["..."],
      "cong_dan_so": ["..."]
    }}
  }},
  "do_dung_day_hoc": {{
    "giao_vien": ["..."],
    "hoc_sinh": ["..."]
  }},
  "tien_trinh_day_hoc": [
    {{
      "ten_hoat_dong": "Hoạt động 1 - Khởi động",
      "thoi_gian": "5 phút",
      "giao_vien": ["..."],
      "hoc_sinh": ["..."]
    }},
    {{
      "ten_hoat_dong": "Hoạt động 2 - Hình thành kiến thức mới",
      "thoi_gian": "15 phút",
      "giao_vien": ["..."],
      "hoc_sinh": ["..."]
    }},
    {{
      "ten_hoat_dong": "Hoạt động 3 - Thực hành - luyện tập",
      "thoi_gian": "10 phút",
      "giao_vien": ["..."],
      "hoc_sinh": ["..."]
    }},
    {{
      "ten_hoat_dong": "Hoạt động 4 - Vận dụng",
      "thoi_gian": "5 phút",
      "giao_vien": ["..."],
      "hoc_sinh": ["..."]
    }}
  ],
  "dieu_chinh_sau_tiet_day": ["..."]
}}
"""


# =========================
# 3. CSS GIAO DIỆN
# =========================
st.markdown("""
<style>
    [data-testid="stAppViewContainer"] {
        background-color: #f4f6f9;
    }

    .main-header {
        background: linear-gradient(135deg, #004e92 0%, #000428 100%);
        padding: 30px;
        border-radius: 15px;
        text-align: center;
        color: white !important;
        margin-bottom: 30px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }

    .main-header h1 {
        color: white !important;
        margin: 0;
        font-size: 2rem;
    }

    .main-header p {
        color: #e0e0e0 !important;
        margin-top: 10px;
        font-style: italic;
    }

    .section-header {
        color: #004e92;
        border-bottom: 2px solid #ddd;
        padding-bottom: 5px;
        margin-top: 20px;
        margin-bottom: 15px;
        font-weight: bold;
    }

    .lesson-plan-paper {
        background-color: white;
        padding: 40px;
        border-radius: 5px;
        border: 1px solid #ccc;
        box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        font-family: 'Times New Roman', Times, serif !important;
        font-size: 14pt !important;
        line-height: 1.6 !important;
        color: #000000 !important;
        text-align: justify;
        white-space: normal;
        word-wrap: break-word;
    }

    div.stButton > button {
        background: linear-gradient(90deg, #11998e, #38ef7d);
        color: white !important;
        border: none;
        padding: 15px 30px;
        font-weight: bold;
        border-radius: 10px;
        width: 100%;
        margin-top: 10px;
        font-size: 18px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.2);
    }

    div.stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(0,0,0,0.3);
    }
</style>
""", unsafe_allow_html=True)


# =========================
# 4. HEADER
# =========================
st.markdown(f"""
<div class="main-header">
    <h1>{TEN_UNG_DUNG}</h1>
    <p>Tác giả: {TEN_TAC_GIA} - {TEN_TRUONG} - ĐT: {SO_DIEN_THOAI}</p>
</div>
""", unsafe_allow_html=True)


# =========================
# 5. CẤU HÌNH API KEY
# =========================
api_key = ""

if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    with st.sidebar:
        st.header("🔐 Cấu hình")
        api_key = st.text_input("Nhập API Key:", type="password")

client = None
if api_key:
    try:
        client = genai.Client(api_key=api_key)
    except Exception as e:
        st.error(f"Lỗi khởi tạo Gemini client: {e}")


# =========================
# 6. TÀI LIỆU NGUỒN
# =========================
st.markdown('<div class="section-header">📂 1. TÀI LIỆU NGUỒN</div>', unsafe_allow_html=True)

has_framework = False
if os.path.exists(FILE_KHUNG_NANG_LUC):
    st.success(f"✅ Đã tự động tích hợp: {FILE_KHUNG_NANG_LUC}")
    has_framework = True
else:
    st.info(f"ℹ️ Chưa có file '{FILE_KHUNG_NANG_LUC}'. Có thể upload thêm để dùng tính năng Năng lực số.")

uploaded_files = st.file_uploader(
    "Tải Ảnh/PDF bài dạy (Kéo thả vào đây):",
    type=["jpg", "jpeg", "png", "pdf"],
    accept_multiple_files=True
)

if uploaded_files:
    st.caption("👁️ Xem trước tài liệu:")
    cols = st.columns(3)
    for i, f in enumerate(uploaded_files):
        if f.type in ["image/jpeg", "image/png"]:
            with cols[i % 3]:
                st.image(f, caption=f.name)
        else:
            with cols[i % 3]:
                st.info(f"📄 {f.name}")


# =========================
# 7. THÔNG TIN BÀI DẠY
# =========================
st.markdown('<div class="section-header">📝 2. THÔNG TIN BÀI DẠY</div>', unsafe_allow_html=True)

c1, c2 = st.columns(2)
with c1:
    lop = st.text_input("📚 Lớp:", "Lớp 4")
with c2:
    ten_bai = st.text_input("📌 Tên bài học:", placeholder="Ví dụ: Học hát bài...")

noidung_bosung = st.text_area("✍️ Ghi chú thêm (nội dung/kiến thức):", height=100)
yeu_cau_them = st.text_input("💡 Yêu cầu đặc biệt:", placeholder="Ví dụ: Tích hợp trò chơi khởi động...")

st.markdown("<br>", unsafe_allow_html=True)


# =========================
# 8. XỬ LÝ CHÍNH
# =========================
if st.button("🚀 SOẠN GIÁO ÁN NGAY"):
    if not api_key:
        st.error("Thiếu API Key.")
    elif not client:
        st.error("Không khởi tạo được Gemini client.")
    elif not uploaded_files and not noidung_bosung and not has_framework:
        st.warning("Thiếu tài liệu đầu vào.")
    elif not ten_bai.strip():
        st.warning("Vui lòng nhập tên bài học.")
    else:
        temp_paths = []
        uploaded_refs = []

        try:
            with st.spinner("AI đang soạn giáo án theo đúng biểu mẫu..."):
                prompt_instruction = generate_lesson_prompt(
                    ten_bai=ten_bai,
                    lop=lop,
                    noidung_bosung=noidung_bosung,
                    yeu_cau_them=yeu_cau_them
                )

                if has_framework and os.path.exists(FILE_KHUNG_NANG_LUC):
                    uploaded_refs.append(
                        upload_file_to_gemini(client, FILE_KHUNG_NANG_LUC, "application/pdf")
                    )

                if uploaded_files:
                    for f in uploaded_files:
                        suffix = os.path.splitext(f.name)[1] or ".tmp"
                        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                            tmp.write(f.getvalue())
                            tmp_path = tmp.name
                            temp_paths.append(tmp_path)

                        mime_type = f.type if f.type else "application/octet-stream"
                        uploaded_refs.append(upload_file_to_gemini(client, tmp_path, mime_type))

                contents = build_contents(
                    prompt_instruction=prompt_instruction,
                    uploaded_refs=uploaded_refs,
                    extra_text=noidung_bosung if noidung_bosung else None
                )

                response = client.models.generate_content(
                    model=MODEL_NAME,
                    contents=contents,
                    config=types.GenerateContentConfig(
                        temperature=0.2,
                        top_p=0.9,
                        max_output_tokens=8192
                    )
                )

                result_text = getattr(response, "text", None)
                if not result_text:
                    raise ValueError("Model không trả về nội dung.")

                lesson_plan_data = extract_json_block(result_text)
                lesson_plan_data = ensure_default_adjustments(lesson_plan_data)
                validate_lesson_plan(lesson_plan_data)

                preview_text = render_lesson_plan_text(lesson_plan_data)
                render_response_box(preview_text)

                doc = create_doc_from_json(lesson_plan_data)
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)

                safe_file_name = re.sub(r'[\\\\/*?:"<>|]', "_", ten_bai).strip() or "GiaoAn"

                st.download_button(
                    label="⬇️ TẢI FILE WORD CHUẨN A4",
                    data=buf,
                    file_name=f"GiaoAn_{safe_file_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

        except Exception as e:
            st.error(f"Có lỗi xảy ra: {e}")

        finally:
            for p in temp_paths:
                try:
                    if os.path.exists(p):
                        os.remove(p)
                except Exception:
                    pass


# =========================
# 9. FOOTER
# =========================
st.markdown("---")
st.markdown(
    f"<div style='text-align: center; color: #666;'>© 2025 - {TEN_TAC_GIA} - {TEN_TRUONG} - ĐT: {SO_DIEN_THOAI}</div>",
    unsafe_allow_html=True
)